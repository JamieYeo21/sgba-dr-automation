# BSF_writer_headline.py
"""
Created on Monday 20 January 2025, 14:11:13

Author: Harry Simmons
"""

import docx
from docx.shared import Pt, Cm, RGBColor
from docx.enum.table import WD_ROW_HEIGHT_RULE, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml.ns import qn
import sys
import os

# Add the Utility folder to sys.path
folder_path = os.path.abspath(os.path.join(os.getcwd(), '..', 'Utility'))  # Replace 'folder_name' with the folder's name
sys.path.append(folder_path)

from BSF.BSF_variables import BSF_variable_creator
from Utility.functions import Change_line_in_DR, format_percentage, create_table

def BSF_headline_writer(BSF_headline_dict, dates_variables, DR):
    cutoff = dates_variables['cutoff']
    last_month = dates_variables['last_month']
    this_month = dates_variables['this_month']

    BSF_BSF_5_total = BSF_headline_dict['BSF_BSF_5_total']
    BSF_started_nc_no = BSF_headline_dict['BSF_started_nc_no']
    BSF_started_nc_pct = BSF_headline_dict['BSF_started_nc_pct']
    BSF_signoff_c_no = BSF_headline_dict['BSF_signoff_c_no']
    BSF_signoff_c_pct = BSF_headline_dict['BSF_signoff_c_pct']
    BSF_started_c_no = BSF_headline_dict['BSF_started_c_no']
    BSF_started_c_pct = BSF_headline_dict['BSF_started_c_pct']
    BSF_started_c_line = BSF_headline_dict['BSF_started_c_line']
    BSF_signoff_c_line = BSF_headline_dict['BSF_signoff_c_line']

    # Headline Title
    text = f'Building Safety Fund (BSF) – monthly update (as at end {this_month}) since previous publication.'
    paragraph = DR.add_paragraph(text, style = 'Heading 3')

    # Paragraph
    text = f'As at {cutoff}, of the {BSF_BSF_5_total} high-rise (18 metres and over in height) residential buildings proceeding with an application for funding through the Building Safety Fund, {BSF_started_nc_no} buildings ({BSF_started_nc_pct}) have started remediation works and {BSF_signoff_c_no} buildings ({BSF_signoff_c_pct}) have completed remediation on unsafe non-ACM cladding, including those awaiting building control sign-off.'
    paragraph = DR.add_paragraph(text, style = 'Normal')

    # Paragraph
    text = f'Overall, {BSF_started_c_no} high-rise buildings ({BSF_started_c_pct}) in the BSF have either started or completed remediation works on non-ACM cladding, {BSF_started_c_line} since the end of {last_month}. Of these, {BSF_signoff_c_no} buildings ({BSF_signoff_c_pct} of buildings) have completed remediation works, {BSF_signoff_c_line} since the end of {last_month}.'
    paragraph = DR.add_paragraph(text, style = 'Normal')