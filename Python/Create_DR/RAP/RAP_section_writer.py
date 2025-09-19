# RAP_writer_section.py
"""
Created on Tuesday 26 August 2025, 13:42:36

Author: Matthew Bandura
"""

import docx
from docx.shared import Pt, Cm, RGBColor
from docx.enum.table import WD_ROW_HEIGHT_RULE, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml.ns import qn
import sys
import os

# Add the Utility folder to sys.path
folder_path = os.path.abspath(os.path.join(os.getcwd(), '..', 'Utility'))  # Replace 'folder_name' with the folder's name
sys.path.append(folder_path)

from RAP.RAP_variables import RAP_variable_creator
from Utility.functions import format_percentage, add_hyperlink
import Utility.docx_svg_patch


def RAP_section_writer(RAP_section_dict,figure_count, dates_variables, paths_variables, DR):
#    class Namespace:
#        def __init__(self, **kwargs):
#            self.__dict__.update(kwargs)

#    rap = Namespace(**RAP_section_dict)
#    print(rap.RAP_18m_complete_no) 

    
    # Unpacking date variables
    dev_cutoff = dates_variables['dev_cutoff']
    cutoff = dates_variables['cutoff']
    this_quarter = dates_variables['end_quarter_word']
    figure_path = os.path.join(paths_variables['figure_path'], f'Figure{figure_count}.svg')

    RAP_18m_complete_no = RAP_section_dict['RAP_18m_complete_no']
    RAP_18m_underway_no = RAP_section_dict['RAP_18m_underway_no']
    RAP_18m_programme_no = RAP_section_dict['RAP_18m_programme_no']

    RAP_18m_est_complete_high_pct = RAP_section_dict['RAP_18m_est_complete_high_pct']
    RAP_18m_est_complete_low_pct = RAP_section_dict['RAP_18m_est_complete_low_pct']

    RAP_18m_est_underway_high_pct = RAP_section_dict['RAP_18m_est_underway_high_pct']
    RAP_18m_est_underway_low_pct = RAP_section_dict['RAP_18m_est_underway_low_pct']

    RAP_11m_est_complete_high_pct = RAP_section_dict['RAP_11m_est_complete_high_pct']
    RAP_11m_est_complete_low_pct = RAP_section_dict['RAP_11m_est_complete_low_pct']

    RAP_11m_total_complete_no = RAP_section_dict['RAP_11m_total_complete_no']
    RAP_11m_total_programme_no = RAP_section_dict['RAP_11m_total_programme_no']

    RAP_11m_est_programme_high_pct = RAP_section_dict['RAP_11m_est_programme_high_pct']
    RAP_11m_est_programme_low_pct = RAP_section_dict['RAP_11m_est_programme_low_pct']

    RAP_11m_est_remaining_low_no = RAP_section_dict['RAP_11m_est_remaining_low_no']
    RAP_11m_est_remaining_high_no = RAP_section_dict['RAP_11m_est_remaining_high_no']

    RAP_11m_est_remaining_low_pct = RAP_section_dict['RAP_11m_est_remaining_low_pct']
    RAP_11m_est_remaining_high_pct = RAP_section_dict['RAP_11m_est_remaining_high_pct']

    # Section Title 
    paragraph = DR.add_paragraph('Remediation Acceleration Plan', style = 'Heading 2')

    # Intro
    paragraph = DR.add_paragraph('MHCLG\'s ', style = 'Normal')  
    add_hyperlink(paragraph, 'Remediation Acceleration Plan ', 'https://www.gov.uk/government/publications/accelerating-remediation-a-plan-for-increasing-the-pace-of-remediation-of-buildings-with-unsafe-cladding-in-england')
    paragraph.add_run('and its ')
    add_hyperlink(paragraph, 'update,', 'https://www.gov.uk/government/publications/remediation-acceleration-plan-update-july-2025')
    paragraph.add_run(' set out targets for the remediation of unsafe cladding on 11m+ buildings.')

    # Paragraph 1
    paragraph = DR.add_paragraph(style = 'Normal')
    run = paragraph.add_run('By the end of 2029, every 18m+ residential building in a government funded scheme will be remediated.')
    run.bold = True

    #Paragraph 2
    paragraph = DR.add_paragraph(f'As at {cutoff}, {RAP_18m_complete_no} 18m+ buildings in a government funded scheme, an estimated {RAP_18m_est_complete_high_pct}-{RAP_18m_est_complete_low_pct} of 18m+ buildings expected to be remediated in a government funded scheme, have completed remediation. A further {RAP_18m_underway_no} buildings, {RAP_18m_est_underway_high_pct}-{RAP_18m_est_underway_low_pct}, have remediation works underway.', style = 'Normal')

    #Paragraph 3
    paragraph = DR.add_paragraph('The estimates of the number of buildings to be remediated in a government funded scheme are based on funding eligibility criteria as of January 2025. These estimates will be updated to reflect the latest funding eligibility criteria.')

    # Figure Title
    paragraph = DR.add_paragraph(style = 'Normal')
    text = f'Figure {figure_count}: {RAP_18m_complete_no} 18m+ buildings in government funded schemes have completed remediation works on unsafe cladding, and a further {RAP_18m_underway_no} buildings have remediation works underway.'
    run = paragraph.add_run(text)
    run.bold = True
    
    # Figure
    DR.add_picture(figure_path, width=Cm(17))
    figure_count += 1

    # Paragraph 4
    paragraph = DR.add_paragraph(style = 'Normal')
    run = paragraph.add_run('By the end of 2029, every 11m+ building with unsafe cladding will either have been remediated, have a date for completion, or its landlords will be liable for penalties.')
    run.bold = True

    # Paragraph 5
    text = f'As at {cutoff}, {RAP_11m_total_complete_no} 11m+ buildings, an estimated {RAP_11m_est_complete_high_pct}-{RAP_11m_est_complete_low_pct} of 11m+ buildings expected to be remediated in the department’s remediation programmes, have completed remediation. A further {RAP_11m_total_programme_no} buildings, {RAP_11m_est_programme_high_pct}-{RAP_11m_est_programme_low_pct}, are already in a remediation programme but are yet to complete remediation works, so are on track to meet the target. '
    DR.add_paragraph(text, style = 'Normal')

    # Paragraph 6
    text = f'There are a remaining {RAP_11m_est_remaining_low_no}-{RAP_11m_est_remaining_high_no} estimated to have unsafe cladding yet to be brought into one of the department’s remediation programmes, {RAP_11m_est_remaining_low_pct}-{RAP_11m_est_remaining_high_pct} of the estimated buildings to be remediated in one of the department’s remediation programmes.'
    DR.add_paragraph(text, style = 'Normal')

    # Figure Title
    paragraph = DR.add_paragraph(style = 'Normal')
    text = f'Figure {figure_count}: {RAP_11m_total_complete_no} 11m+ buildings have completed remediation works on unsafe cladding, a further {RAP_11m_total_programme_no} buildings are in a remediation programme, and an estimated {RAP_11m_est_remaining_low_no}-{RAP_11m_est_remaining_high_no} buildings are yet to be brought into a remediation programme.'
    run = paragraph.add_run(text)
    run.bold = True
    
    # Figure
    DR.add_picture(figure_path, width=Cm(17))
    figure_count += 1

    # Heading 
    paragraph = DR.add_paragraph('Additional stretch targets', style = 'Heading 3')

    # Paragraph 8
    paragraph = DR.add_paragraph('To date, 39 developers have signed a ', style = 'Normal')  
    add_hyperlink(paragraph, 'joint plan', 'https://www.gov.uk/government/publications/joint-plan-to-accelerate-developer-led-remediation-and-improve-resident-experience')
    paragraph.add_run(f' with the government which includes remediation stretch targets. The latest data, as of {dev_cutoff}, on developers’ progress against these stretch targets are published in the ')
    add_hyperlink(paragraph, 'Remediation Acceleration Plan Update', 'https://www.gov.uk/government/publications/remediation-acceleration-plan-update-july-2025')
    paragraph.add_run(f'. Updated data, as at {cutoff}, will be published in {this_quarter}.')


    # Paragraph 9
    text = 'To date, 113 registered providers of social housing have signed a joint plan with the government which includes remediation stretch targets. Further data on the progress in meeting these targets will be published in the future.'
    DR.add_paragraph(text, style = 'Normal')


    return figure_count