# CSS_writer_headline.py
"""
Created on Friday 17 January 2025, 14:33:41

Author: Harry Simmons
"""

import docx
from docx.shared import Pt, Cm, RGBColor
from docx.enum.table import WD_ROW_HEIGHT_RULE, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml.ns import qn
import sys
import os

# Add the Utility folder to sys.path
folder_path = os.path.abspath(os.path.join(os.getcwd(), '..', 'Utility'))  # Replace 'folder_name' with the folder's name
sys.path.append(folder_path)

from Utility.functions import Change_line_in_DR, format_percentage, create_table

def CSS_headline_writer(CSS_headline_dict, dates_variables, DR):
    cutoff = dates_variables['cutoff']
    last_month = dates_variables['last_month']
    this_month = dates_variables['this_month']

    CSS_eligible_total = CSS_headline_dict['CSS_eligible_total']
    CSS_BSF_transfer = CSS_headline_dict['CSS_BSF_transfer']
    CSS_eligible_total_line = CSS_headline_dict['CSS_eligible_total_line']
    CSS_started_c_no = CSS_headline_dict['CSS_started_c_no']
    CSS_started_c_pct = CSS_headline_dict['CSS_started_c_pct']  # Assuming you have this variable
    CSS_started_c_line = CSS_headline_dict['CSS_started_c_line']
    CSS_completed_nc_no = CSS_headline_dict['CSS_completed_nc_no']
    CSS_completed_c_pct = CSS_headline_dict['CSS_completed_c_pct']  # Assuming you have this variable
    CSS_completed_c_line = CSS_headline_dict['CSS_completed_c_line']
    CSS_pre_eligible_total = CSS_headline_dict['CSS_pre_eligible_total']
    CSS_pre_eligible = CSS_headline_dict['CSS_pre_eligible']
    CSS_pre_application = CSS_headline_dict['CSS_pre_application']

    # Headline Title
    text = f'Cladding Safety Scheme (CSS) – monthly update (as at end {this_month}) since previous publication.'
    paragraph = DR.add_paragraph(text, style = 'Heading 3')

    # Paragraph
    text = f'As at {cutoff}, {CSS_eligible_total} buildings 11 metres and over in height have been assessed as eligible for the Cladding Safety Scheme (including {CSS_BSF_transfer} buildings that have transferred from the BSF), {CSS_eligible_total_line} since the end of {last_month}. Of these, {CSS_started_c_no} buildings ({CSS_started_c_pct}) have either started or completed remediation works, {CSS_started_c_line} since the end of {last_month}. Of these, {CSS_completed_nc_no} buildings ({CSS_completed_c_pct}) have completed remediation works, including those awaiting building control sign-off, {CSS_completed_c_line} since the end of {last_month}.'
    paragraph = DR.add_paragraph(text, style = 'Normal')

    # Paragraph
    text = f'The CSS continues to investigate and pull in potentially eligible buildings. There are a further {CSS_pre_eligible_total} buildings 11 metres and over in height in the pre-eligible stages of the Cladding Safety Scheme, which launched fully in July 2023. Of these, {CSS_pre_eligible} buildings are progressing through eligibility checks, and {CSS_pre_application} buildings are in the pre-application stage.'
    paragraph = DR.add_paragraph(text, style = 'Normal')