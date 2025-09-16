# DR_start_infrastructure.py
"""
Created on Thursday 13 March 2025, 10:05:24

Author: Harry Simmons
"""

import docx
from docx.shared import Pt, Cm, RGBColor
import sys
import os

def DR_building_safety_overview(DR):
    # Section Title 
    paragraph = DR.add_paragraph('Building Safety Overview', style = 'Heading 2')

    # Paragraph
    DR.add_paragraph('This data release publishes data across all government remediation activities to give an overview of the status of progress to remediate unsafe cladding on residential buildings over 11m in England. This includes:')

    # Bullet point
    DR.add_paragraph('Estimates of the number of buildings that have or had unsafe cladding to be remediated in a government remediation programme.', style = 'List Bullet')

    # Bullet point
    DR.add_paragraph('Data relating to the Aluminium Composite Material (ACM) cladding programme – those highest risk buildings that are high-rise buildings with unsafe, ‘Grenfell-style’ ACM cladding.', style = 'List Bullet')

    # Bullet point
    DR.add_paragraph('Data relating to the Building Safety Fund, which funds the remediation of eligible high-rise residential buildings with other forms of unsafe cladding.', style = 'List Bullet')

    # Bullet point
    DR.add_paragraph('Data relating to the Cladding Safety Scheme, which funds the remediation of residential buildings over 11m in height with unsafe cladding.', style = 'List Bullet')

    # Bullet point
    DR.add_paragraph('Data relating to developer-led remediation, which reports on those buildings that developers have agreed to remediate.', style = 'List Bullet')

    # Bullet point
    DR.add_paragraph('Data on residential buildings over 11m in height that are the responsibility of registered social housing providers.', style = 'List Bullet')