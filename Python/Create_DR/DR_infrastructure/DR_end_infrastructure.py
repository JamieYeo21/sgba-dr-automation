# DR_start_infrastructure.py
"""
Created on Thursday 13 March 2025, 10:05:24

Author: Harry Simmons
"""

import docx
from docx.shared import Pt, Cm, RGBColor
import sys
import os

folder_path = os.path.abspath(os.path.join(os.getcwd(), '..', 'Utility'))  # Replace 'folder_name' with the folder's name
sys.path.append(folder_path)

from Utility.functions import add_hyperlink

def DR_end(DR, dates_variables):

    this_month = dates_variables['this_month']

    #converts this month into a format which will work in the hyperlink
    hyperlink_month = dates_variables['hyperlink_month']
    
    # Section Title 
    paragraph = DR.add_paragraph('Accompanying dashboard', style = 'Heading 2')

    # Paragraph
    paragraph = DR.add_paragraph('An additional interactive dashboard [INSERT LINK] showing the information in this release is available.', style = 'Normal')


    # Section Title 
    paragraph = DR.add_paragraph('Accompanying tables', style = 'Heading 2')

    # Paragraph
    paragraph = DR.add_paragraph('Additional ', style = 'Normal')
    add_hyperlink(paragraph, 'management information tables', f'https://www.gov.uk/government/publications/building-safety-remediation-monthly-data-release-{hyperlink_month}')
    paragraph.add_run(' are available.')

    # Paragraph
    DR.add_paragraph('The tables provide data on:', style = 'Normal')

    # Bullet point
    DR.add_paragraph('the remediation progress of high-rise (18 metres and over) residential buildings identified with unsafe Aluminium Composite Material (ACM) cladding systems,', style = 'List Bullet')
    
    # Bullet point
    DR.add_paragraph('the remediation progress of high-rise residential buildings with unsafe non-ACM cladding systems that are pursuing successful applications from their Building Safety Fund (BSF) Registration,', style = 'List Bullet')

    # Bullet point
    DR.add_paragraph('data on buildings in the Cladding Safety Scheme (CSS),', style = 'List Bullet')

    # Bullet point
    DR.add_paragraph('the remediation progress of buildings covered by the developer remediation contract, including a developer-by-developer breakdown,', style = 'List Bullet')

    # Bullet point
    DR.add_paragraph('the remediation progress of buildings monitored under the social housing survey, including a provider-by-provider breakdown,', style = 'List Bullet')

    # Bullet point
    DR.add_paragraph('the progress of the Waking Watch Relief Fund and Waking Watch Replacement Fund, and', style = 'List Bullet')

    # Bullet point
    DR.add_paragraph('building safety enforcement action undertaken by Local Authorities in England.', style = 'List Bullet')

    # Section Title 
    paragraph = DR.add_paragraph('Related Statistics', style = 'Heading 2')

    # Heading
    paragraph = DR.add_paragraph('BRE Testing', style = 'Heading 3')

    # Paragraph
    paragraph = DR.add_paragraph("Previously, MHCLG published a ", style = 'Normal')
    add_hyperlink(paragraph, "table on samples", "https://gbr01.safelinks.protection.outlook.com/?url=https%3A%2F%2Fassets.publishing.service.gov.uk%2Fmedia%2F5d9de664ed915d35d0dcca58%2FTable_4_Building_Safety_Data_Release_September_2019.csv&data=05%7C02%7CHolly.Turner%40levellingup.gov.uk%7Ca3beedc540d64b9b07e208dc7b1c15e4%7Cbf3468109c7d43dea87224a2ef3995a8%7C0%7C0%7C638520607758745355%7CUnknown%7CTWFpbGZsb3d8eyJWIjoiMC4wLjAwMDAiLCJQIjoiV2luMzIiLCJBTiI6Ik1haWwiLCJXVCI6Mn0%3D%7C0%7C%7C%7C&sdata=9Ou2NuiGcT1nn5yD1YRLNvMGxWaPGDETmm4F8x5vEUk%3D&reserved=0")
    paragraph.add_run(' received by BRE for testing which has been discontinued as of October 2019 (see ')
    add_hyperlink(paragraph, 'see technical notes document',  f'https://www.gov.uk/government/publications/building-safety-remediation-monthly-data-release-july-2025')
    paragraph.add_run('). The ')
    add_hyperlink(paragraph, "data table of descriptions of large-scale system tests", "https://gbr01.safelinks.protection.outlook.com/?url=https%3A%2F%2Fassets.publishing.service.gov.uk%2Fmedia%2F5fac156ce90e075c51b736da%2FTable_4_Building_Safety_Data_Release_October_2020.csv&data=05%7C02%7CHolly.Turner%40levellingup.gov.uk%7Ca3beedc540d64b9b07e208dc7b1c15e4%7Cbf3468109c7d43dea87224a2ef3995a8%7C0%7C0%7C638520607758758036%7CUnknown%7CTWFpbGZsb3d8eyJWIjoiMC4wLjAwMDAiLCJQIjoiV2luMzIiLCJBTiI6Ik1haWwiLCJXVCI6Mn0%3D%7C0%7C%7C%7C&sdata=cOTllZJQFFJvcleJKXTgxUyjUj9KmAGe%2B07ACvlJx50%3D&reserved=0")
    paragraph.add_run(" undertaken by the BRE and the number of buildings with similar cladding systems was discontinued in November 2020.")

    # Heading
    paragraph = DR.add_paragraph('Developer Data', style = 'Heading 3')

    # Paragraph
    paragraph = DR.add_paragraph(style = 'Normal')
    paragraph.add_run('MHCLG has published data ')
    add_hyperlink(paragraph, 'provided by developers', 'https://www.gov.uk/government/publications/building-safety-remediation-monthly-data-release-may-2025')
    paragraph.add_run(' who have signed the developer remediation contract. This release provides information on the number of buildings in scope of the contract, assessments in place, number of buildings requiring remediation works and status of those works by developer.')

    # Heading
    paragraph = DR.add_paragraph('English Housing Survey: Feeling Safe from Fire', style = 'Heading 3')

    # Paragraph
    paragraph = DR.add_paragraph("MHCLG has published the ", style = 'Normal')
    add_hyperlink(paragraph, "English Housing Survey 2020 to 2021: Feeling Safe from Fire report", "https://www.gov.uk/government/statistics/english-housing-survey-2020-to-2021-feeling-safe-from-fire")
    paragraph.add_run(", providing information on the extent to which people feel safe from fire in their homes.")

    # Heading
    paragraph = DR.add_paragraph('Estimating the prevalence and costs of external wall system life-safety fire risk in mid-rise residential buildings', style = 'Heading 3')

    # Paragraph
    paragraph = DR.add_paragraph("MHCLG has published data on the ", style = 'Normal')
    add_hyperlink(paragraph, "prevalence of external wall system life-safety fire risk in mid-rise (11-18m) residential buildings in England", "https://www.gov.uk/government/publications/estimating-the-prevalence-and-costs-of-external-wall-system-life-safety-fire-risk-in-mid-rise-residential-buildings-in-england")
    paragraph.add_run(", and the estimated cost as at July 2021 to remediate or mitigate these buildings. On 17th July, MHCLG published an ")
    add_hyperlink(paragraph, "updated estimate of the prevalence of external wall system fire risk in mid-rise buildings", "https://www.gov.uk/government/publications/building-safety-remediation-monthly-data-release-june-2025")
    paragraph.add_run(". Should these figures change further, MHCLG will publish a new update.")

    # Heading
    paragraph = DR.add_paragraph('EWS1 requirements on residential buildings in England', style = 'Heading 3')

    # Paragraph
    paragraph = DR.add_paragraph("MHCLG publishes ", style = 'Normal')
    add_hyperlink(paragraph, "quarterly data on the numbers of EWS1 forms", "https://www.gov.uk/government/collections/ews1-or-equivalent-lender-data-on-mortgage-valuations-for-flats")
    paragraph.add_run(" (or equivalent) that have been required on mortgage valuations for flats.")

    # Heading
    paragraph = DR.add_paragraph('Population and Dwelling Numbers', style = 'Heading 3')

    # Paragraph
    paragraph = DR.add_paragraph("Previously, MHCLG published estimates on population and dwelling numbers of residential buildings in the ", style = 'Normal')
    add_hyperlink(paragraph, "Building Safety Programme data release", "https://www.gov.uk/government/publications/building-safety-programme-monthly-data-release-september-2023")
    paragraph.add_run(". Should these figures change, MHCLG will publish a new update. On 17 July, MHCLG published an ")
    add_hyperlink(paragraph, "updated estimate of the number of 11-18m residential buildings in England", "https://www.gov.uk/government/publications/building-safety-remediation-monthly-data-release-june-2025")
    paragraph.add_run(". Should these figures change further, MHCLG will publish a new update.")


    # Heading
    paragraph = DR.add_paragraph('RSH publication', style = 'Heading 3')

    # Paragraph
    paragraph = DR.add_paragraph("On 20 March 2025, the Regulator of Social Housing published ", style = 'Normal')
    add_hyperlink(paragraph, "findings from the Fire Safety Remediation Survey (FRS)", "https://www.gov.uk/government/collections/fire-safety-remediation-in-social-housing-in-england")
    paragraph.add_run(" for buildings 11 metres and over in height as at 31 December 2024, which opened to all landlords on 13 December 2024 and closed on 22 January 2025.")

    # Heading
    paragraph = DR.add_paragraph('Social Housing Provider Data', style = 'Heading 3')

    # Paragraph
    paragraph = DR.add_paragraph(style = 'Normal')
    run = paragraph.add_run('MHCLG has published ')
    run = paragraph.add_run('data provided by social housing providers')
    run = paragraph.add_run(' on remediation progress of their building stock. This release provides information on the number of buildings, assessments in place, number of buildings requiring remediation works and status of those works by social housing provider.')

    # Heading
    paragraph = DR.add_paragraph('Waking Watch costs', style = 'Heading 3')

    # Paragraph
    paragraph = DR.add_paragraph("On 16 October 2020, MHCLG published ", style = 'Normal')
    add_hyperlink(paragraph, "information on Waking Watch costs", "https://www.gov.uk/government/publications/building-safety-programme-waking-watch-costs")
    paragraph.add_run(" based on data collected through a range of external stakeholders from July to September 2020.")

    # Heading
    paragraph = DR.add_paragraph('Cladding remediation unit costs', style = 'Heading 3')

    # Paragraph
    paragraph = DR.add_paragraph("On 17 December 2024, MHCLG published data on ", style = 'Normal')
    add_hyperlink(paragraph, "cladding remediation unit costs", "https://www.gov.uk/government/publications/cladding-remediation-unit-costs-analysis-of-high-rise-non-acm-buildings")
    paragraph.add_run(", providing data on costs per square metre of cladding remediated for high-rise non-ACM buildings, including analysis by cladding area and cost categories.")

    # Section Title 
    paragraph = DR.add_paragraph('Technical note', style = 'Heading 2')

    # Paragraph
    paragraph = DR.add_paragraph(style = 'Normal')
    paragraph.add_run('Please see the accompanying ')
    add_hyperlink(paragraph,'technical notes document', f'https://www.gov.uk/government/publications/building-safety-remediation-monthly-data-release-{hyperlink_month}')
    paragraph.add_run(' document for further details.')
