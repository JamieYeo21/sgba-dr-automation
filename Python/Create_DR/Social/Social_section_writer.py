# Social_writer_section.py
"""
Created on Thursday 29 May 2025, 14:38:00

Author: Harry Simmons
"""

from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.shared import Pt, Cm, RGBColor


def Social_section_writer(DR, figure_count):
    paragraph = DR.add_paragraph(style = 'Heading 1')
    run = paragraph.add_run('[ADD SOCIAL SECTION HERE]')
    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    run.bold = True

    # Figure Title
    paragraph = DR.add_paragraph(style = 'Normal')
    text = f'Figure {figure_count}: '

    run = paragraph.add_run(text)
    run.bold = True
    run = paragraph.add_run('[INSERT STAT HERE]')
    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    text = ' of social buildings identified to have unsafe cladding have started or completed remediation works, with '
    run = paragraph.add_run(text)
    run.bold = True
    run = paragraph.add_run('[INSERT STAT HERE]')
    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    text = ' (of identified buildings) having completed remediation works.'
    run = paragraph.add_run(text)
    run.bold = True

    # Figure
    DR.add_picture(f'Q:\BSP\Automation\DR Automation\DR_outputs\DR_graphs\Figure{figure_count}.svg', width=Cm(17))
    figure_count += 1
    

    # Figure Title
    paragraph = DR.add_paragraph(style = 'Normal')
    text = f'Figure {figure_count}: '
    run = paragraph.add_run(text)
    run.bold = True
    run = paragraph.add_run('[INSERT STAT HERE]')
    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    text = ' of the 18m+ social buildings identified to have unsafe cladding have started or completed remediation, compared to '
    run = paragraph.add_run(text)
    run.bold = True
    run = paragraph.add_run('[INSERT STAT HERE]')
    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    text = ' of the 11-18m buildings.'
    run = paragraph.add_run(text)
    run.bold = True

    # Figure
    DR.add_picture(f'Q:\BSP\Automation\DR Automation\DR_outputs\DR_graphs\Figure{figure_count}.svg', width=Cm(17))
    figure_count += 1