# Publishing_Tool.py
"""
Created on Thursday 17 April 2025, 15:11:00

Author: Harry Simmons
"""

import os
import re
from docx import Document
from docx.oxml.ns import qn
from zipfile import ZipFile


print('RUNNING...')


def get_word_path(folder_path):
    files = os.listdir(folder_path)
    word_files = [file for file in files if file.endswith('.docx') or file.endswith('.doc')]
    if len(word_files) == 1:
        word_path = os.path.join(folder_path, word_files[0])
        return word_path
    else:
        raise FileNotFoundError("No Word document found in the folder.")
    

def extract_month_year_from_filename(filename):
    # Remove extension
    name_without_ext = os.path.splitext(filename)[0]

    # Define regex for month and year
    month_pattern = r'\b(January|February|March|April|May|June|July|August|September|October|November|December)\b'
    year_pattern = r'\b(20\d{2})\b'

    month_match = re.search(month_pattern, name_without_ext, re.IGNORECASE)
    year_match = re.search(year_pattern, name_without_ext)

    if month_match and year_match:
        month = month_match.group(1).capitalize()
        year = year_match.group(1)
        return month, year
    else:
        month = 'NO'
        year = 'DATE'
        return month, year


def extract_headings_to_md(input_folder, output_folder):
    word_path = get_word_path(input_folder)
    filename = os.path.basename(word_path)
    month, year = extract_month_year_from_filename(filename)

    # Create dynamic markdown path
    md_filename = f"publishing_markdown_{month}_{year}.md"
    md_path = os.path.join(output_folder, md_filename)
    doc = Document(word_path)
    md_lines = []

    rels = doc.part.rels

    for element in doc.element.body:
        if element.tag == qn('w:p'): # Paragraph
            para = next(p for p in doc.paragraphs if p._element == element)
            if para.style.name.startswith('Heading'):
                level = int(para.style.name.split()[-1])
                md_lines.append(f"{'#' * level} {para.text}")
            elif para.style.name in ['List Bullet', 'List Number']:
                md_lines.append(f"* {para.text}")         
            else:
                para_text = ""     
                for run in para.runs:
                    text = run.text
                    text = re.sub(r'(\b[\w\.-]+@[\w\.-]+\.\w{2,}\b)', r'<\1>', text)
                    parent = run._element.getparent()
                    if parent.tag == qn('w:hyperlink'):
                        r_id = parent.get(qn('r:id'))
                        if r_id and r_id in rels:
                            url = rels[r_id].target_ref
                            para_text += f"Text: {text} -> URL: {url}"
                        else:
                            para_text += text
                    elif run.bold:
                        para_text += f"**{text}**"
                    else:
                        para_text += text
                md_lines.append(para_text)
        elif element.tag == qn('w:tbl'): # Table
            table = next(t for t in doc.tables if t._element == element)
            rows = table.rows
            if not rows:
                continue
            header = [cell.text.strip() for cell in rows[0].cells]
            md_lines.append('|'.join(header))
            md_lines.append('|'.join(['---'] + ['---:' for _ in header[1:]]))
            for row in rows[1:]:
                row_data = [cell.text.strip() for cell in row.cells]
                md_lines.append('|'.join(row_data))

    with open(md_path, 'w', encoding='utf-8') as md_file:
        md_file.write('\n'.join(md_lines))

extract_headings_to_md(
    'Q:\\BSP\Automation\\DR Automation\\Excel_inputs\\[PUT DR HERE]',
    'Q:\\BSP\Automation\\DR Automation\\DR_outputs\\DR_markdown'
)


print('DONE!')